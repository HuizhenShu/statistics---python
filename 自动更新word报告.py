
# coding: utf-8

# In[1]:


import docx
import xlwt
import xlrd
doc=docx.Document(r'测试报告-8小时保湿功效测试-第1版-20180809模板.docx')
len(doc.paragraphs)


# In[2]:


def pro_code(excelfile):
    workbook = xlrd.open_workbook(excelfile)
    sheet_names= workbook.sheet_names()
    sheet = workbook.sheet_by_name('产品对照')
    nrows = sheet.nrows
    prodict = {}
    for i in range(nrows):
        prodict[sheet.cell_value(i,0)] = sheet.cell_value(i,1)
    #print(prodict)
    return prodict
pro_code(u'保湿模板所需数据1.xlsx')
prodict = pro_code(u'保湿模板所需数据1.xlsx')
pro = ['A','B','C','D']
#print(prodict.keys())


# In[3]:


def excel_word(excelfile,name,docfile,outfile,n):
    doc=docx.Document(docfile)
    #print(len(doc.paragraphs))
#     for i in doc.picture:
#         print(i)
#替换文本
#     for i in doc.paragraphs:
#         #print(i.text)
#         for j in range(len(pro)):
#             if '~' in i.text:
#                 #print(i.text)
#                 i.text = i.text.replace('%s~'%pro[j],prodict[pro[j]])
                #print(i.text)
#     print()
#插入表格
    workbook = xlrd.open_workbook(excelfile)
    sheet_names= workbook.sheet_names()
    #print(sheet_names)
    sheet2 = workbook.sheet_by_name(name)
    nrows = sheet2.nrows    # 获取行总数
    ncols = sheet2.ncols    #获取列总数
    #table = doc.add_table(rows=nrows, cols=ncols)#在word中创建表格
    print(ncols,nrows)
#     for table in doc.tables:
#         print(doc.tables)
    table =  doc.tables[n]
#         for row in table.rows:  # 遍历表格的所有行
#             for cell in row.cells:
    for i in range(ncols):
        for j in range(nrows):
#             print (sheet2.cell_value(j,i))# 获取第四行内容
            #print(table.rows[j].cells[i].text)
            table.rows[j].cells[i].text = sheet2.cell_value(j, i)
            if sheet2.cell_value(0, i) == 'PRODUCT_CODE':
                table.rows[j].cells[i].text = prodict[sheet2.cell_value(j, i)]
            #table.rows[j].cells[i].text = sheet2.cell_value(j, i)
    doc.save(outfile)
   
#excel_word(u'C181808 保湿模板所需数据.xlsx','Sheet1',r'C181808 测试报告模板.docx',r'C181808 测试报告结果.docx',0)
#excel_word(u'C181808 保湿模板所需数据.xlsx','Sheet2',r'C181808 测试报告模板.docx',r'C181808 测试报告结果.docx',1)


# In[4]:


columns = ['PRODUCT_CODE','对比时间点','N','t值','P值','显著性']
#将显著性表格插入
def excel_word2(excelfile,name,docfile,outfile,n):
    doc=docx.Document(docfile)
    #print(len(doc.paragraphs))
#     for i in doc.picture:
#         print(i)
#     for i in doc.paragraphs:
#         print(i.text)
#     print()
    workbook = xlrd.open_workbook(excelfile)
    sheet_names= workbook.sheet_names()
    #print(sheet_names)
    sheet2 = workbook.sheet_by_name(name)
    nrows = sheet2.nrows    # 获取行总数
    ncols = sheet2.ncols    #获取列总数
    #table = doc.add_table(rows=nrows, cols=ncols)#在word中创建表格
    print(ncols,nrows)
#     for table in doc.tables:
#         print(doc.tables)
    table =  doc.tables[n]
    print(len(table.rows))
#         for row in table.rows:  # 遍历表格的所有行
#             for cell in row.cells:
#     for i in range(ncols):
    for j in range(nrows):
#             print (sheet2.cell_value(j,i))# 获取第四行内容
#         print(table.rows[j].cells[2].text)
#         print(table.rows[j].cells[3].text)
#         table.rows[j].cells[3].text = sheet2.cell_value(j, 0)
        for i in range(ncols):
           # print(sheet2.cell_value(0,i))
#             if sheet2.cell_value(0, i) == columns[i-1]:
            print(sheet2.cell_value(j,i))
            table.rows[j].cells[i].text =str(sheet2.cell_value(j, i))
            
            if  sheet2.cell_value(0, i) == 'PRODUCT_CODE':
                #print(sheet2.cell_value(j, i))
                table.rows[j].cells[0].text = prodict[(sheet2.cell_value(j, i))]
#             elif  sheet2.cell_value(0, i) == 'plist':
#                 table.rows[j].cells[3].text = sheet2.cell_value(j, i)         
        #table.rows[j].cells[0].text = sheet2.cell_value(j, 1)
    doc.save(outfile)


# In[6]:


excel_word(u'output.xlsx','不同产品时间均值标准误',r'测试报告-8小时保湿功效测试-第1版-20180809模板.docx',r'测试报告-8小时保湿功效测试-第1版-20180809.docx',2)
excel_word2(u'output.xlsx','不同产品时间均值差异p值',r'测试报告-8小时保湿功效测试-第1版-20180809.docx',r'测试报告-8小时保湿功效测试-第1版-20180809.docx',3)
excel_word(u'output.xlsx','不同产品时间差均值标准误',r'测试报告-8小时保湿功效测试-第1版-20180809.docx',r'测试报告-8小时保湿功效测试-第1版-20180809.docx',4)
excel_word2(u'output.xlsx','不同产品时间差均值差异p值',r'测试报告-8小时保湿功效测试-第1版-20180809.docx',r'测试报告-8小时保湿功效测试-第1版-20180809.docx',5)


# In[ ]:


import re
lista = ['和六神登飞来峰','斤斤计较']
ff = '会计分录圣诞节发生率'
aa = re.match('.*圣诞.*',ff).group()
#aa = re.match('com','comwww.runcomoob')
print(aa)
def word_in_p(listp,word):
    for i in listp:
        try:
            aa = re.match('.*%s.*'%word,i).group()
            print(lista.index(aa))
        except:
            pass
word_in_p(lista,'和')
#print(lista.index(re.))

